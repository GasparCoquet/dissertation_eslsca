"""
build_dissertation.py
Assemble the cleaned/honest section markdown files into a single dissertation
in BOTH Word (.docx) and PDF. Lightweight markdown subset: #..#### headings,
paragraphs with **bold**/*italic*, "- " bullets, "> " quotes, pipe tables,
and ![caption](path) images.
"""
import os, re, sys, string
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from fpdf import FPDF
from fpdf.enums import XPos, YPos

ROOT = os.path.dirname(os.path.dirname(__file__))
# Section markdown sources: CLI arg overrides; default to the in-repo folder.
SCRATCH = sys.argv[1] if len(sys.argv) > 1 else os.path.join(ROOT, "manuscript", "sections")
DOCX_OUT = os.path.join(ROOT, "Dissertation_Coquet_FINAL.docx")
PDF_OUT = os.path.join(ROOT, "Dissertation_Coquet_FINAL.pdf")

SECTIONS = ["sec_abstract.md", "sec_frontmatter.md", "sec_intro_clean.md",
            "sec_ch1_clean.md", "sec_ch2_clean.md", "sec_ch3_clean.md",
            "sec_ch4_clean.md", "sec_conclusion_clean.md", "sec_references_clean.md",
            "sec_appendixA.md", "sec_appendixB.md", "sec_appendixC.md", "sec_appendixD.md"]

# ---------------------------------------------------------------- markdown parse
def parse(md):
    lines = md.replace("\r", "").split("\n")
    blocks, i, n = [], 0, len(lines)
    while i < n:
        ln = lines[i]
        if not ln.strip():
            i += 1; continue
        if ln.strip().startswith("```"):           # fenced code / display equation
            i += 1; code = []
            while i < n and not lines[i].strip().startswith("```"):
                code.append(lines[i]); i += 1
            i += 1                                  # consume closing fence
            blocks.append(("code", code)); continue
        m = re.match(r'^(#{1,4})\s+(.*)', ln)
        if m:
            blocks.append(("h", len(m.group(1)), m.group(2).strip())); i += 1; continue
        m = re.match(r'^!\[(.*?)\]\((.*?)\)', ln)
        if m:
            blocks.append(("img", m.group(2).strip(), m.group(1).strip())); i += 1; continue
        if ln.lstrip().startswith("|"):
            tb = []
            while i < n and lines[i].lstrip().startswith("|"):
                tb.append(lines[i]); i += 1
            rows = []
            for r in tb:
                cells = [c.strip() for c in r.strip().strip("|").split("|")]
                if all(re.match(r'^:?-{2,}:?$', c) for c in cells):  # separator
                    continue
                rows.append(cells)
            if rows:
                blocks.append(("table", rows))
            continue
        if ln.lstrip().startswith("- "):
            items = []
            while i < n and lines[i].lstrip().startswith("- "):
                items.append(lines[i].lstrip()[2:].strip()); i += 1
            blocks.append(("ul", items)); continue
        if ln.lstrip().startswith(">"):
            q = []
            while i < n and lines[i].lstrip().startswith(">"):
                q.append(lines[i].lstrip().lstrip(">").strip()); i += 1
            blocks.append(("quote", " ".join(x for x in q if x))); continue
        # paragraph
        para = []
        while i < n and lines[i].strip() and not re.match(r'^(#{1,4}\s|!\[|\||- |>)', lines[i].lstrip()):
            para.append(lines[i].strip()); i += 1
        blocks.append(("p", " ".join(para)))
    return blocks

_PUNCT_SET = set(string.punctuation)
def _split_emph(text):
    """Tokenise -> list of (text, bold, italic). Handles **bold** and *italic*
    using CommonMark-style flanking, so math stars (V*, Q*, lambda*, n*) are
    left LITERAL rather than being mistaken for italic delimiters and stripped."""
    text = text.replace("`", "")          # drop inline-code backticks
    n = len(text)
    def can_open(i, run):
        nxt = text[i + run] if i + run < n else " "
        prv = text[i - 1] if i - 1 >= 0 else " "
        if nxt.isspace():
            return False
        if nxt in _PUNCT_SET and not (prv.isspace() or prv in _PUNCT_SET):
            return False
        return True
    def can_close(j, run):
        prv = text[j - 1] if j - 1 >= 0 else " "
        nxt = text[j + run] if j + run < n else " "
        if prv.isspace():
            return False
        if prv in _PUNCT_SET and not (nxt.isspace() or nxt in _PUNCT_SET):
            return False
        return True
    tokens, buf, i = [], [], 0
    while i < n:
        if text[i] == "*":
            run = 2 if text[i:i + 2] == "**" else 1
            if can_open(i, run):
                j, found = i + run, -1
                while j < n:
                    if text[j] == "*":
                        crun = 2 if text[j:j + 2] == "**" else 1
                        if crun >= run and j > i + run and can_close(j, run):
                            found = j; break
                        j += crun; continue
                    j += 1
                if found >= 0:
                    if buf:
                        tokens.append(("".join(buf), False, False)); buf = []
                    tokens.append((text[i + run:found], run == 2, run == 1))
                    i = found + run; continue
            buf.append(text[i:i + run]); i += run; continue
        buf.append(text[i]); i += 1
    if buf:
        tokens.append(("".join(buf), False, False))
    return tokens or [(text, False, False)]

def inline(text):
    """-> list of (text, bold, italic) for the DOCX path."""
    return _split_emph(text)

_PUNCT = {"—": ", ", "–": "-", "‘": "'", "’": "'",
          "“": '"', "”": '"', "…": "...", " ": " "}
def plain(text):
    """Strip markdown emphasis AND normalise typographic punctuation to glyphs
    that render cleanly in the PDF. Used only on the PDF path (DOCX uses inline)."""
    text = "".join(t for t, _b, _i in _split_emph(text))
    for k, v in _PUNCT.items():
        text = text.replace(k, v)
    return text

# ---------------------------------------------------------------- load sections
def load_blocks():
    allb = []
    for s in SECTIONS:
        allb.append(("section_break", s))
        allb += parse(open(os.path.join(SCRATCH, s), encoding="utf-8").read())
    return allb

EQ_DPI = 200          # equation PNGs are rendered at this dpi (see render_equations.py)
def png_size(path):
    """Return (width_px, height_px) from a PNG header without extra deps."""
    import struct
    with open(path, "rb") as f:
        head = f.read(24)
    return struct.unpack(">II", head[16:24])

def is_equation_img(rel_path):
    return os.path.basename(rel_path).startswith("eq_")

CHAPTER_STARTS = ("chapter", "general introduction", "general conclusion",
                  "references", "appendix", "abstract", "acknowledg",
                  "declaration", "list of")
def is_chapter_h(level, text):
    return level == 1 and text.lower().startswith(CHAPTER_STARTS)

# ---------------------------------------------------------------- DOCX
def set_base_style(doc):
    # School spec: Times New Roman, 12pt body, 1.5 line spacing, justified, 2.5cm margins
    st = doc.styles["Normal"]; st.font.name = "Times New Roman"; st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = st.paragraph_format; pf.line_spacing = 1.5; pf.space_after = Pt(6)
    HSIZE = {1: 14, 2: 14, 3: 13, 4: 12}
    for lvl, sz in HSIZE.items():
        try:
            hs = doc.styles[f"Heading {lvl}"]
            hs.font.name = "Times New Roman"; hs.font.size = Pt(sz)
            hs.font.color.rgb = RGBColor(0, 0, 0); hs.font.bold = True
            hs.element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except KeyError:
            pass

def add_toc_field(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), r'TOC \o "1-3" \h \z \u')
    t = OxmlElement('w:t'); t.text = "Right-click here and 'Update Field' to build the table of contents."
    r = OxmlElement('w:r'); r.append(t); fld.append(r); p._p.append(fld)

def build_docx(blocks):
    doc = Document(); set_base_style(doc)
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    # Title page
    def center(txt, size, bold=False, italic=False, space=6):
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.size = Pt(size); r.bold = bold; r.italic = italic
        p.paragraph_format.space_after = Pt(space); return p
    for _ in range(3): doc.add_paragraph()
    center("ESLSCA BUSINESS SCHOOL PARIS", 13, bold=True)
    center("MBA in Trading & Financial Markets", 12)
    for _ in range(2): doc.add_paragraph()
    center("To what extent does Reinforcement Learning reduce market impact "
           "costs (Implementation Shortfall) on fragmented European equity markets?",
           17, bold=True, space=10)
    for _ in range(2): doc.add_paragraph()
    center("Research Dissertation", 13, italic=True)
    center("Gaspar COQUET", 13, bold=True)
    center("Dissertation Supervisor: Abdelbaric EL ALAOUI", 11.5)
    center("Academic Year 2025–2026", 11.5)
    for _ in range(2): doc.add_paragraph()
    center("All empirical results in this dissertation are simulated Implementation "
           "Shortfall under a transparent, literature-calibrated market model. No "
           "proprietary or real tick data is used; only signs, rankings and comparative "
           "statics are interpreted. The full pipeline is reproducible from the "
           "accompanying code repository.", 9.5, italic=True)
    doc.add_page_break()
    # TOC
    h = doc.add_paragraph(); r = h.add_run("Table of Contents"); r.bold = True; r.font.size = Pt(15)
    add_toc_field(doc); doc.add_page_break()

    for b in blocks:
        kind = b[0]
        if kind == "section_break":
            continue
        if kind == "h":
            _, level, text = b
            if is_chapter_h(level, text):
                doc.add_page_break()
            p = doc.add_heading(text, level=min(level, 4))
            for r in p.runs: r.font.color.rgb = RGBColor(0, 0, 0)
        elif kind == "p":
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            for tx, bold, ital in inline(b[1]):
                r = p.add_run(tx); r.bold = bold; r.italic = ital
        elif kind == "ul":
            for it in b[1]:
                p = doc.add_paragraph(style="List Bullet")
                p.paragraph_format.line_spacing = 1.5      # match body (school spec)
                for tx, bold, ital in inline(it):
                    r = p.add_run(tx); r.bold = bold; r.italic = ital
        elif kind == "code":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.left_indent = Inches(0.3)
            lines_ = b[1] or [""]
            for j, cl in enumerate(lines_):
                r = p.add_run(cl)
                r.font.name = "Consolas"; r.font.size = Pt(10.5)
                r.element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
                if j < len(lines_) - 1:
                    r.add_break()
        elif kind == "quote":
            p = doc.add_paragraph(); p.paragraph_format.left_indent = Inches(0.4)
            p.paragraph_format.line_spacing = 1.5
            for tx, bold, ital in inline(b[1]):
                r = p.add_run(tx); r.italic = True; r.bold = bold
        elif kind == "table":
            rows = b[1]; ncol = max(len(r) for r in rows)
            t = doc.add_table(rows=0, cols=ncol); t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            for ri, row in enumerate(rows):
                cells = t.add_row().cells
                for ci in range(ncol):
                    txt = row[ci] if ci < len(row) else ""
                    cell = cells[ci]; cell.paragraphs[0].text = ""
                    for tx, bold, ital in inline(txt):
                        run = cell.paragraphs[0].add_run(tx)
                        run.bold = bold or (ri == 0); run.italic = ital
                        run.font.size = Pt(9)
        elif kind == "img":
            path = os.path.join(ROOT, b[1])
            if os.path.exists(path):
                if is_equation_img(b[1]):
                    wpx, _ = png_size(path)
                    w_in = min(6.3, wpx / EQ_DPI)         # natural size, capped to text width
                    doc.add_picture(path, width=Inches(w_in))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if b[2]:                               # optional equation label/caption
                        cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        rc = cap.add_run(b[2]); rc.italic = True; rc.font.size = Pt(9)
                else:
                    doc.add_picture(path, width=Inches(5.7))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    rc = cap.add_run(b[2]); rc.italic = True; rc.font.size = Pt(9)
    doc.save(DOCX_OUT)
    print("DOCX written:", DOCX_OUT)

# ---------------------------------------------------------------- PDF
FONT = "TimesNR"        # body/headings: Times New Roman (school spec). NOT "Times":
                        # that is a reserved PDF core-font name (Latin-1 only) and
                        # add_font would be silently ignored, dropping Greek/math glyphs.
DEJA = "DejaVu"         # glyph-coverage fallback for Greek/math Times may lack
MONO = "DejaVuMono"
def register_fonts(pdf):
    import matplotlib
    fdir = os.path.join(matplotlib.get_data_path(), "fonts", "ttf")
    win = r"C:\Windows\Fonts"
    times = {"": "times.ttf", "B": "timesbd.ttf", "I": "timesi.ttf"}
    if all(os.path.exists(os.path.join(win, f)) for f in times.values()):
        for style, f in times.items():
            pdf.add_font(FONT, style, os.path.join(win, f))
    else:                # Times not installed: fall back to DejaVu serif as body
        pdf.add_font(FONT, "", os.path.join(fdir, "DejaVuSerif.ttf"))
        pdf.add_font(FONT, "B", os.path.join(fdir, "DejaVuSerif-Bold.ttf"))
        pdf.add_font(FONT, "I", os.path.join(fdir, "DejaVuSerif-Italic.ttf"))
    # Register DejaVu serif as a fallback so Greek/math glyphs missing from
    # Times New Roman (sigma, lambda, in, sqrt, Sigma, ...) still render.
    pdf.add_font(DEJA, "", os.path.join(fdir, "DejaVuSerif.ttf"))
    pdf.add_font(DEJA, "B", os.path.join(fdir, "DejaVuSerif-Bold.ttf"))
    pdf.add_font(DEJA, "I", os.path.join(fdir, "DejaVuSerif-Italic.ttf"))
    pdf.add_font(MONO, "", os.path.join(fdir, "DejaVuSansMono.ttf"))
    try:
        pdf.set_fallback_fonts([DEJA])
    except Exception:
        pass

class PDF(FPDF):
    toc_entries = []
    _safe_mc = True
    def multi_cell(self, *a, **kw):
        if self._safe_mc and "new_x" not in kw:
            kw["new_x"] = XPos.LMARGIN; kw["new_y"] = YPos.NEXT
        return super().multi_cell(*a, **kw)
    def header(self):
        if self.page_no() <= self.front_pages: return
        self.set_font(FONT, "I", 8); self.set_text_color(120)
        self.cell(0, 8, "Reinforcement Learning and Implementation Shortfall, G. Coquet, ESLSCA MBA 2025-2026", align="C")
        self.set_text_color(0); self.ln(9)
    def footer(self):
        if self.page_no() <= self.front_pages: return
        self.set_y(-14); self.set_font(FONT, "", 9); self.set_text_color(120)
        self.cell(0, 10, str(self.page_no() - self.front_pages), align="C"); self.set_text_color(0)

def render_pdf_body(pdf, blocks, record=False):
    for b in blocks:
        kind = b[0]
        if kind == "section_break":
            continue
        if kind == "h":
            _, level, text = b
            if is_chapter_h(level, text):
                pdf.add_page()
            sizes = {1: 14, 2: 14, 3: 13, 4: 12}
            if record and level <= 2:
                pdf.toc_entries.append((level, text, pdf.page_no()))
            pdf.ln(2.5); pdf.set_font(FONT, "B", sizes.get(level, 12))
            pdf.multi_cell(0, 7.2, plain(text)); pdf.ln(1.5)
        elif kind == "p":
            pdf.set_font(FONT, "", 12)
            pdf.multi_cell(0, 6.4, plain(b[1]), align="J"); pdf.ln(2)
        elif kind == "ul":
            pdf.set_font(FONT, "", 12)
            for it in b[1]:
                pdf.multi_cell(0, 6.2, "•  " + plain(it));
            pdf.ln(2)
        elif kind == "code":
            pdf.set_font(MONO, "", 9.5); pdf.ln(0.5)
            for cl in (b[1] or [""]):
                pdf.set_x(pdf.l_margin + 4)
                pdf.multi_cell(0, 5.0, cl.rstrip())
            pdf.set_font(FONT, "", 12); pdf.ln(2)
        elif kind == "quote":
            pdf.set_font(FONT, "I", 12)
            pdf.set_x(pdf.l_margin + 6)
            pdf.multi_cell(0, 6.2, plain(b[1])); pdf.ln(1.5)
        elif kind == "table":
            render_pdf_table(pdf, b[1])
        elif kind == "img":
            path = os.path.join(ROOT, b[1])
            if os.path.exists(path):
                if is_equation_img(b[1]):
                    wpx, hpx = png_size(path)
                    w_mm = min(160.0, wpx / EQ_DPI * 25.4)   # natural size, capped
                    h_mm = w_mm * hpx / wpx
                    y0 = pdf.get_y()
                    if y0 + h_mm + 8 > pdf.h - pdf.b_margin:
                        pdf.add_page(); y0 = pdf.get_y()
                    pdf.image(path, x=(210 - w_mm) / 2, y=y0 + 1.5, w=w_mm)
                    pdf.set_y(y0 + 1.5 + h_mm + 3)           # advance below the equation
                else:
                    if pdf.get_y() > 200: pdf.add_page()
                    w = 150; pdf.image(path, x=(210 - w) / 2, w=w)
                    pdf.set_font(FONT, "I", 9); pdf.set_text_color(90)
                    pdf.multi_cell(0, 5, plain(b[2]), align="C"); pdf.set_text_color(0); pdf.ln(2)

def render_pdf_table(pdf, rows):
    ncol = max(len(r) for r in rows)
    norm = [[plain(row[ci]) if ci < len(row) else "" for ci in range(ncol)] for row in rows]
    pdf.set_font(FONT, "", 8.5); pdf.ln(1)
    pdf._safe_mc = False
    try:
        with pdf.table(text_align="CENTER", first_row_as_headings=True,
                       borders_layout="ALL", line_height=4.8) as table:
            for row in norm:
                tr = table.row()
                for c in row:
                    tr.cell(c)
    finally:
        pdf._safe_mc = True
    pdf.ln(2)

def build_pdf(blocks):
    pdf = PDF(format="A4"); pdf.front_pages = 0; register_fonts(pdf)
    pdf.set_margins(25, 25, 25); pdf.set_auto_page_break(True, 25)
    # Pass 1: title+abstract page count + body to collect TOC entries & numbers
    # Front matter
    def front(pdf):
        pdf.add_page(); pdf.ln(28)
        pdf.set_font(FONT, "B", 13); pdf.multi_cell(0, 7, "ESLSCA BUSINESS SCHOOL PARIS", align="C")
        pdf.set_font(FONT, "", 12); pdf.multi_cell(0, 7, "MBA in Trading & Financial Markets", align="C"); pdf.ln(8)
        pdf.set_font(FONT, "B", 17)
        pdf.multi_cell(0, 9, "To what extent does Reinforcement Learning reduce market impact "
                             "costs (Implementation Shortfall) on fragmented European equity markets?", align="C")
        pdf.ln(8); pdf.set_font(FONT, "I", 13); pdf.multi_cell(0, 7, "Research Dissertation", align="C")
        pdf.set_font(FONT, "B", 13); pdf.multi_cell(0, 7, "Gaspar COQUET", align="C")
        pdf.set_font(FONT, "", 11.5)
        pdf.multi_cell(0, 6, "Dissertation Supervisor: Abdelbaric EL ALAOUI", align="C")
        pdf.multi_cell(0, 6, "Academic Year 2025-2026", align="C"); pdf.ln(10)
        pdf.set_font(FONT, "I", 9.5); pdf.set_text_color(90)
        pdf.multi_cell(0, 5, "All empirical results in this dissertation are simulated Implementation Shortfall "
            "under a transparent, literature-calibrated market model. No proprietary or real tick data is used; "
            "only signs, rankings and comparative statics are interpreted. The full pipeline is reproducible "
            "from the accompanying code repository.", align="C")
        pdf.set_text_color(0)
    # ---- pass 1 (throwaway) to get heading page numbers ----
    pdf.toc_entries = []
    front(pdf)
    n_toc_pages = 2
    for _ in range(n_toc_pages): pdf.add_page()
    pdf.front_pages = 1 + n_toc_pages
    body_start = pdf.page_no()
    render_pdf_body(pdf, blocks, record=True)
    toc = pdf.toc_entries

    # ---- pass 2 (final) ----
    pdf2 = PDF(format="A4"); pdf2.front_pages = 1 + n_toc_pages; register_fonts(pdf2)
    pdf2.set_margins(25, 25, 25); pdf2.set_auto_page_break(True, 25)
    pdf2.toc_entries = []
    front(pdf2)
    # TOC pages
    pdf2.add_page()
    pdf2.set_font(FONT, "B", 15); pdf2.multi_cell(0, 9, "Table of Contents"); pdf2.ln(2)
    for level, text, pg in toc:
        disp = pg - pdf.front_pages  # printed page number
        pdf2.set_font(FONT, "B" if level == 1 else "", 11 if level == 1 else 10)
        indent = "" if level == 1 else "    "
        title = plain(text)
        # dotted leader
        line = f"{indent}{title}"
        pdf2.cell(0, 6, line, align="L")
        pdf2.cell(0, 6, str(disp), align="R", new_x="LMARGIN", new_y="NEXT")
    # pad to n_toc_pages
    while pdf2.page_no() < 1 + n_toc_pages:
        pdf2.add_page()
    render_pdf_body(pdf2, blocks, record=False)
    pdf2.output(PDF_OUT)
    print("PDF written:", PDF_OUT, "pages:", pdf2.page_no())

if __name__ == "__main__":
    blocks = load_blocks()
    build_docx(blocks)
    build_pdf(blocks)
